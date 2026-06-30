"""
Service de génération du compte rendu au format .docx, fidèle au gabarit
"COMPTE RENDU COMITE D'ETUDE" d'Afriland First Bank.

Principe : on part du fichier Word original (app/templates/template_compte_rendu.docx),
qui contient déjà l'en-tête, le pied de page, le logo, la page de garde et la table
des matières (champ Word). On repère les sections par leur titre et les tableaux par
leur position dans le document, puis on duplique une ligne "modèle" pour chaque
donnée réelle de la réunion. Le fichier d'origine n'est jamais modifié sur disque :
on travaille sur une copie chargée en mémoire.
"""

import copy
import io
import json
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

# Constantes de mise en forme alignées sur le gabarit Afriland First Bank
FONT_NAME = "Century Gothic"
RED = RGBColor(0xFF, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK_FILL = "000000"

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "templates", "template_compte_rendu.docx"
)


# ---------------------------------------------------------------------------
# Utilitaires bas niveau
# ---------------------------------------------------------------------------

def _iter_block_items(parent):
    """Parcourt les paragraphes ET les tableaux du corps du document, dans
    l'ordre où ils apparaissent (recette standard python-docx)."""
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _is_toc_paragraph(paragraph):
    """La table des matières répète les intitulés de chaque section : il faut
    l'ignorer pour ne pas confondre une entrée de sommaire avec le vrai titre
    de section plus bas dans le document."""
    return paragraph.style.name.lower().startswith("toc")


def _is_section_heading(paragraph):
    """Détecte un titre de section (Heading 2 / Titre2 selon la langue de
    Word). Sert à savoir où une section S'ARRÊTE, pour ne jamais aller
    chercher un tableau ou un paragraphe appartenant à la section SUIVANTE."""
    name = paragraph.style.name.lower()
    return name.startswith("heading") or name.startswith("titre")


def _table_after_heading(doc, heading_text):
    """Retourne le 1er tableau qui suit IMMÉDIATEMENT (sans dépasser le titre
    de section suivant) un paragraphe contenant heading_text. Retourne None
    si la section ne contient pas de tableau (ex: contenu en texte libre) —
    ne JAMAIS continuer à chercher au-delà, sous peine de capturer/corrompre
    le tableau d'une autre section plus loin dans le document."""
    waiting = False
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            if _is_toc_paragraph(block):
                continue
            text = block.text.strip()
            if waiting and text and _is_section_heading(block):
                return None
            if heading_text.lower() in text.lower():
                waiting = True
        elif isinstance(block, Table) and waiting:
            return block
    return None


def _paragraph_after_heading(doc, heading_text):
    """Retourne le 1er paragraphe non vide qui suit IMMÉDIATEMENT un
    paragraphe contenant heading_text (utilisé pour les sections texte libre
    : veille, divers, résolutions en liste...). Ne dépasse jamais le titre de
    section suivant."""
    waiting = False
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            if _is_toc_paragraph(block):
                continue
            text = block.text.strip()
            if waiting and text:
                if _is_section_heading(block):
                    return None
                return block
            if heading_text.lower() in text.lower():
                waiting = True
        elif waiting:
            # Un tableau apparaît avant tout paragraphe texte : cette section
            # n'est pas une liste de paragraphes (cf. _fill_resolutions).
            return None
    return None


def _set_paragraph_text(paragraph, text):
    """Réécrit le texte d'un paragraphe en conservant le style du 1er run."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _set_cell_text(cell, text):
    """Réécrit le texte d'une cellule (1ère ligne) en gardant sa mise en forme."""
    paragraph = cell.paragraphs[0]
    _set_paragraph_text(paragraph, text)
    # Supprime les paragraphes additionnels éventuels dans la cellule
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    # Si la cellule est vidée, on retire aussi la puce de numérotation :
    # sinon une puce isolée ("-") reste visible dans une cellule "vide".
    if not text:
        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                pPr.remove(numPr)


def _take_template_row(table, template_row_index):
    """Récupère une copie indépendante (deepcopy) de la ligne-modèle AVANT
    toute suppression de lignes, pour pouvoir la cloner ensuite autant de
    fois que nécessaire."""
    return copy.deepcopy(table.rows[template_row_index]._tr)


def _append_cloned_row(table, template_tr):
    """Ajoute en fin de tableau un nouveau clone de la ligne-modèle capturée
    via _take_template_row, et retourne la nouvelle ligne."""
    new_tr = copy.deepcopy(template_tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _clear_data_rows(table, header_row_count):
    """Supprime toutes les lignes de données existantes (gabarit d'exemple),
    en ne gardant que les header_row_count premières lignes (en-têtes)."""
    rows_to_remove = table.rows[header_row_count:]
    for row in rows_to_remove:
        row._tr.getparent().remove(row._tr)


def _find_textbox_paragraph(doc, contains_text):
    """Cherche un paragraphe dont le texte contient contains_text, y compris
    DANS les zones de texte (textboxes) de la page de garde, que python-docx
    n'expose pas via doc.paragraphs."""
    body = doc.element.body
    for p_element in body.iter(qn("w:p")):
        paragraph = Paragraph(p_element, doc)
        if contains_text.lower() in paragraph.text.lower():
            return paragraph
    return None


def _load_json_list(raw):
    """Parse un champ JSON stocké en base (Text). Retourne toujours une liste."""
    if not raw:
        return []
    try:
        data = json.loads(raw)
        return data if isinstance(data, list) else []
    except (json.JSONDecodeError, TypeError):
        return []


def _format_run(run, *, font=FONT_NAME, size_pt=None, color=None, bold=None):
    """Applique explicitement police/taille/couleur/gras à un run, plutôt que
    de dépendre uniquement de ce qui a été hérité par copie de la ligne-modèle."""
    if font is not None:
        run.font.name = font
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold


def _format_cell_text(cell, **kwargs):
    """Applique _format_run à tous les runs de la 1ère ligne d'une cellule."""
    for run in cell.paragraphs[0].runs:
        _format_run(run, **kwargs)


def _set_cell_shading(cell, fill_hex):
    """Définit (fill_hex="RRGGBB") ou retire (fill_hex=None) la trame de fond
    d'une cellule, indépendamment de ce que contenait le gabarit d'origine."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if fill_hex is None:
        if shd is not None:
            tcPr.remove(shd)
        return
    if shd is None:
        shd = tcPr.makeelement(qn("w:shd"), {})
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_numero_cell(cell, text):
    """Cellule de numéro de ligne avec fond noir / police blanche en gras
    (style des tableaux ORDRE DU JOUR et RÉCAPITULATIF DES RÉSOLUTIONS)."""
    _set_cell_text(cell, text)
    _set_cell_shading(cell, BLACK_FILL)
    _format_cell_text(cell, size_pt=11, color=WHITE, bold=True)


# ---------------------------------------------------------------------------
# Remplissage des sections
# ---------------------------------------------------------------------------

def _fill_title_page(doc, meeting):
    date_str = meeting.date.strftime("%d/%m/%Y") if meeting.date else ""

    title_p = _find_textbox_paragraph(doc, "COMPTE RENDU COMITE D")
    if title_p is not None:
        _set_paragraph_text(title_p, f"COMPTE RENDU COMITE D’ETUDE DU {date_str}")

    debut_p = _find_textbox_paragraph(doc, "Heure début")
    if debut_p is not None:
        _set_paragraph_text(debut_p, f"Heure début : {meeting.heure_debut or '—'}")

    fin_p = _find_textbox_paragraph(doc, "Heure de fin")
    if fin_p is not None:
        _set_paragraph_text(fin_p, f"Heure de fin : {meeting.heure_fin or '—'}")

    lieu_p = _find_textbox_paragraph(doc, "Lieu :")
    if lieu_p is not None:
        _set_paragraph_text(lieu_p, f"Lieu : {meeting.lieu or '—'}")


def _fill_presence_table(doc, participants):
    table = _table_after_heading(doc, "LISTE DE PRESENCE")
    if table is None:
        return

    presents = [p.name for p in participants if p.present]
    absents = [p.name for p in participants if not p.present]

    template_tr = _take_template_row(table, 1) if len(table.rows) > 1 else None
    _clear_data_rows(table, header_row_count=1)
    _set_cell_text(table.rows[0].cells[0], f"PRESENCES ({len(presents)})")
    _set_cell_text(table.rows[0].cells[1], "ABSENCES")

    if template_tr is None:
        return
    nb_rows = max(len(presents), len(absents), 1)
    for i in range(nb_rows):
        row = _append_cloned_row(table, template_tr)
        _set_cell_text(row.cells[0], presents[i] if i < len(presents) else "")
        _set_cell_text(row.cells[1], absents[i] if i < len(absents) else "")


def _fill_ordre_du_jour(doc, points):
    table = _table_after_heading(doc, "ORDRE DU JOUR")
    if table is None or not table.rows:
        return

    # IMPORTANT : on capture et on vide TOUJOURS, même si `points` est vide.
    # Sinon les points d'exemple du gabarit original (NPL, secteur extractif...)
    # restent affichés tels quels dans le compte rendu généré.
    template_tr = _take_template_row(table, 0)
    _clear_data_rows(table, header_row_count=0)

    if not points:
        row = _append_cloned_row(table, template_tr)
        _set_numero_cell(row.cells[0], "—")
        _set_cell_text(row.cells[1], "Aucun point à l'ordre du jour n'a été identifié.")
        return

    for i, point in enumerate(points, start=1):
        row = _append_cloned_row(table, template_tr)
        _set_numero_cell(row.cells[0], str(i))
        _set_cell_text(row.cells[1], point)


def _is_recap_header_row(row):
    texts = [c.text.strip().upper() for c in row.cells]
    return any("DATE" in t for t in texts) and any("DOSSIER" in t for t in texts)


def _dedupe_identical_header_rows(table, header_row_count):
    """Si plusieurs lignes d'en-tête consécutives ont EXACTEMENT le même texte
    (duplication accidentelle dans le gabarit Word), n'en garde qu'une seule.
    Sans ça, certains lecteurs (Word) peuvent afficher l'en-tête en double."""
    if header_row_count < 2:
        return header_row_count
    first_texts = [c.text for c in table.rows[0].cells]
    keep = 1
    for row in table.rows[1:header_row_count]:
        if [c.text for c in row.cells] == first_texts:
            keep += 1
        else:
            break
    if keep > 1:
        for row in table.rows[1:keep]:
            row._tr.getparent().remove(row._tr)
        return header_row_count - (keep - 1)
    return header_row_count


def _fill_recap_table(doc, activites):
    table = _table_after_heading(
        doc, "Tableau récapitulatif DES activités et recommandations"
    )
    if table is None or not table.rows:
        return

    # IMPORTANT : le nombre de lignes d'en-tête varie selon la version du
    # gabarit (1 ligne dans certains, 2 lignes identiques dupliquées dans
    # d'autres). On le détecte par le CONTENU ("DATE"/"DOSSIER") plutôt que
    # de le deviner via len(table.rows) — une version précédente supposait
    # toujours 1 ligne d'en-tête et, sur les gabarits qui en ont 2, traitait
    # la 2e ligne d'en-tête comme une donnée à vider, faisant disparaître les
    # libellés de colonnes (DATE, DOSSIER, ANALYSTE...).
    header_row_count = 0
    for row in table.rows:
        if _is_recap_header_row(row):
            header_row_count += 1
        else:
            break
    if header_row_count == 0:
        header_row_count = 1  # filet de sécurité si la détection échoue

    # Si le gabarit duplique l'en-tête (2 lignes identiques), on fusionne en 1.
    header_row_count = _dedupe_identical_header_rows(table, header_row_count)

    for header_row in table.rows[:header_row_count]:
        for cell in header_row.cells:
            _set_cell_shading(cell, None)
            _format_cell_text(cell, size_pt=10, color=RED, bold=True)

    if len(table.rows) > header_row_count:
        template_tr = _take_template_row(table, header_row_count)
    else:
        template_tr = _take_template_row(table, header_row_count - 1)
    _clear_data_rows(table, header_row_count=header_row_count)

    if not activites:
        row = _append_cloned_row(table, template_tr)
        _set_cell_text(row.cells[0], "—")
        for col_index in range(1, len(row.cells)):
            _set_cell_text(row.cells[col_index], "")
        _set_cell_text(
            row.cells[2],
            "Aucun dossier d'étude n'a été présenté lors de cette séance.",
        )
        for cell in row.cells:
            _format_cell_text(cell, size_pt=9)
        return

    columns = [
        None,  # n°, géré séparément
        "date",
        "dossier",
        "nature_document",
        "analyste",
        "observations",
        "decision_comite",
        "delais",
        None,  # colonne vide en fin de tableau
    ]

    for i, activite in enumerate(activites, start=1):
        row = _append_cloned_row(table, template_tr)
        _set_cell_text(row.cells[0], str(i))
        for col_index, key in enumerate(columns):
            if key is None:
                continue
            _set_cell_text(row.cells[col_index], str(activite.get(key) or ""))
        # Contenu des lignes de données : taille 9, forcé explicitement.
        for cell in row.cells:
            _format_cell_text(cell, size_pt=9)


def _fill_resolutions_as_table(table, resolutions):
    template_tr = _take_template_row(table, 0)
    _clear_data_rows(table, header_row_count=0)

    if not resolutions:
        row = _append_cloned_row(table, template_tr)
        _set_numero_cell(row.cells[0], "—")
        _set_cell_text(row.cells[1], "Aucune résolution n'a été prise lors de cette séance.")
        return

    for i, resolution in enumerate(resolutions, start=1):
        row = _append_cloned_row(table, template_tr)
        _set_numero_cell(row.cells[0], str(i))
        _set_cell_text(row.cells[1], resolution)


def _fill_resolutions_as_paragraphs(doc, template_paragraph, resolutions):
    if not resolutions:
        _set_paragraph_text(template_paragraph, "RAS")
        return

    _set_paragraph_text(template_paragraph, resolutions[0])
    previous_element = template_paragraph._p
    for resolution in resolutions[1:]:
        new_element = copy.deepcopy(template_paragraph._p)
        previous_element.addnext(new_element)
        _set_paragraph_text(Paragraph(new_element, doc), resolution)
        previous_element = new_element


def _fill_resolutions(doc, resolutions):
    """Remplit la section RÉCAPITULATIF DES RÉSOLUTIONS.

    IMPORTANT : selon la version du gabarit, cette section peut être un vrai
    TABLEAU (numéro à fond noir/police blanche + texte, comme ORDRE DU JOUR)
    OU une simple liste de paragraphes à puces. On détecte dynamiquement la
    structure réellement présente plutôt que d'en supposer une seule — une
    version précédente supposait toujours un tableau et, quand le gabarit
    utilisait des paragraphes, continuait à chercher un tableau plus loin
    dans le document et finissait par corrompre le tableau OPPORTUNITÉS
    D'APPRENDRE rencontré ensuite.
    """
    table = _table_after_heading(doc, "RÉCAPITULATIF DES RÉSOLUTIONS")
    if table is not None and table.rows:
        _fill_resolutions_as_table(table, resolutions)
        return

    paragraph = _paragraph_after_heading(doc, "RÉCAPITULATIF DES RÉSOLUTIONS")
    if paragraph is not None:
        _fill_resolutions_as_paragraphs(doc, paragraph, resolutions)


def _fill_simple_text_section(doc, heading_text, text):
    paragraph = _paragraph_after_heading(doc, heading_text)
    if paragraph is None:
        return
    _set_paragraph_text(paragraph, text or "RAS")


def _fill_opportunites_table(doc, opportunites, solde_tontine):
    table = _table_after_heading(doc, "RÉCAPITULATIF DES OPPORTUNITES D")
    if table is not None and table.rows:
        template_index = 1 if len(table.rows) > 1 else 0
        template_tr = _take_template_row(table, template_index)
        _clear_data_rows(table, header_row_count=1)
        columns = ["nom", "motif", "montant", "date", "statut"]
        for opp in opportunites:
            row = _append_cloned_row(table, template_tr)
            for col_index, key in enumerate(columns):
                _set_cell_text(row.cells[col_index], str(opp.get(key) or ""))

    solde_paragraph = _find_textbox_paragraph(doc, "solde dans la caisse")
    if solde_paragraph is not None and solde_tontine:
        text = solde_paragraph.text
        # On remplace uniquement la partie variable, en gardant la phrase d'origine
        if "……." in text:
            _set_paragraph_text(
                solde_paragraph, text.replace("…….", str(solde_tontine))
            )


def _fill_rapporteurs(doc, rapporteurs_planification, president, rapporteur):
    table = _table_after_heading(doc, "PLANIFICATION DES RAPPORTEURS")
    if table is not None and table.rows:
        # IMPORTANT : on vide TOUJOURS le tableau, même si rapporteurs_planification
        # est vide. C'est le bug principal : auparavant, une liste vide laissait
        # les 11 noms d'exemple du gabarit (ZRA TAKOUA EFRAIM, ANDRE KEVIN...)
        # inchangés -> le même tableau apparaissait sur CHAQUE compte rendu généré,
        # peu importe la réunion réelle.
        template_index = 1 if len(table.rows) > 1 else 0
        template_tr = _take_template_row(table, template_index)
        _clear_data_rows(table, header_row_count=1)
        if rapporteurs_planification:
            for i, nom in enumerate(rapporteurs_planification, start=1):
                row = _append_cloned_row(table, template_tr)
                _set_cell_text(row.cells[0], str(i))
                _set_cell_text(row.cells[1], nom)
        else:
            row = _append_cloned_row(table, template_tr)
            _set_cell_text(row.cells[0], "—")
            _set_cell_text(
                row.cells[1],
                "Aucun rapporteur n'a été planifié lors de cette séance.",
            )

    signature_paragraph = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("Rapporteur") and "Président" in text:
            signature_paragraph = p
            break

    noms_paragraph = None
    if signature_paragraph is not None:
        found_signature = False
        for p in doc.paragraphs:
            if p._p is signature_paragraph._p:
                found_signature = True
                continue
            if found_signature and p.text.strip():
                noms_paragraph = p
                break

    if noms_paragraph is not None and (rapporteur or president):
        _set_paragraph_text(
            noms_paragraph,
            f"             {rapporteur or '—'}                                                 \t\t{president or '—'}",
        )


# ---------------------------------------------------------------------------
# Point d'entrée du service
# ---------------------------------------------------------------------------

def generate_meeting_docx(meeting) -> io.BytesIO:
    """
    Génère le compte rendu .docx d'une réunion, dans le format exact du
    gabarit Afriland First Bank.

    `meeting` est l'objet SQLAlchemy Meeting (avec sa relation .participants
    déjà chargée).
    """
    doc = Document(TEMPLATE_PATH)

    _fill_title_page(doc, meeting)
    _fill_presence_table(doc, meeting.participants)
    _fill_ordre_du_jour(doc, _load_json_list(meeting.ordre_du_jour))
    _fill_recap_table(doc, _load_json_list(meeting.activites_recommandations))
    _fill_resolutions(doc, _load_json_list(meeting.resolutions))
    _fill_simple_text_section(doc, "POINTS SUR LA VEILLE", meeting.points_veille)
    _fill_simple_text_section(doc, "FAITS SAILLANTS", meeting.faits_saillants)
    _fill_simple_text_section(doc, "DIVERS", meeting.divers)
    _fill_opportunites_table(
        doc, _load_json_list(meeting.opportunites_apprendre), meeting.solde_tontine
    )
    _fill_rapporteurs(
        doc,
        _load_json_list(meeting.rapporteurs_planification),
        meeting.president,
        meeting.rapporteur,
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer